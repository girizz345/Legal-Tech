import io
from typing import Sequence
from xml.sax.saxutils import escape

from docx import Document
from docx.shared import Pt
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

from app.modules.ai_orchestration.schemas import SectionDraft


def render_pdf(document_title: str, sections: Sequence[SectionDraft]) -> bytes:
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=A4)
    styles = getSampleStyleSheet()
    note_style = ParagraphStyle(
        "Note", parent=styles["Italic"], textColor="#555555", leftIndent=20, spaceAfter=12
    )

    story = [Paragraph(escape(document_title), styles["Title"]), Spacer(1, 12)]
    for section in sections:
        story.append(Paragraph(escape(section.title), styles["Heading2"]))
        story.append(Paragraph(escape(section.body_text), styles["BodyText"]))
        if section.note_text:
            story.append(Paragraph(f"Note: {escape(section.note_text)}", note_style))
        story.append(Spacer(1, 8))

    doc.build(story)
    return buffer.getvalue()


def render_docx(document_title: str, sections: Sequence[SectionDraft]) -> bytes:
    document = Document()
    document.add_heading(document_title, level=0)
    for section in sections:
        document.add_heading(section.title, level=2)
        document.add_paragraph(section.body_text)
        if section.note_text:
            note_paragraph = document.add_paragraph()
            note_run = note_paragraph.add_run(f"Note: {section.note_text}")
            note_run.italic = True
            note_run.font.size = Pt(9)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
